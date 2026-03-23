import sys
import os

try:
    from docx import Document
except ImportError:
    Document = None

# Mapping from Latin (AAcadHN/Parliament layout) to Unicode Georgian (Mkhedruli)
# This covers standard Georgian QWERTY layout used by legacy fonts like AAcadHN, AcadNusx, Sylfaen (old)
LATIN_TO_GEORGIAN = {
    'a': 'ა', 'b': 'ბ', 'g': 'გ', 'd': 'დ', 'e': 'ე', 'v': 'ვ', 'z': 'ზ', 't': 'თ',
    'i': 'ი', 'k': 'კ', 'l': 'ლ', 'm': 'მ', 'n': 'ნ', 'o': 'ო', 'p': 'პ', 'J': 'ჟ',
    'r': 'რ', 's': 'ს', 'T': 'ტ', 'u': 'უ', 'f': 'ფ', 'q': 'ქ', 'R': 'ღ', 'y': 'ყ',
    'S': 'შ', 'C': 'ჩ', 'c': 'ც', 'Z': 'ძ', 'w': 'წ', 'W': 'ჭ', 'x': 'ხ', 'j': 'ჯ',
    'h': 'ჰ',
    # Common variations or shifts often found in these fonts
    'F': 'თ', # sometimes
    'G': 'ღ', # sometimes
}

def convert_text(text):
    """Converts text from Latin-encoded Georgian to Unicode Georgian."""
    if not text:
        return text

    result = []
    for char in text:
        # Check direct mapping
        if char in LATIN_TO_GEORGIAN:
            result.append(LATIN_TO_GEORGIAN[char])
        else:
            result.append(char)
    return "".join(result)

def process_docx(input_path, output_path):
    if Document is None:
        print("Error: 'python-docx' library is not installed. Please install it using 'pip install python-docx'")
        sys.exit(1)

    doc = Document(input_path)

    # Process paragraphs
    for paragraph in doc.paragraphs:
        # We need to process runs to preserve formatting if possible,
        # but changing text length might mess up runs.
        # Safest is to iterate runs and replace text inside them.
        for run in paragraph.runs:
            run.text = convert_text(run.text)
            # Optional: Force font to Sylfaen
            run.font.name = 'Sylfaen'

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = convert_text(run.text)
                        run.font.name = 'Sylfaen'

    doc.save(output_path)
    print(f"Successfully saved converted document to: {output_path}")

def process_text_file(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    converted = convert_text(content)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(converted)
    print(f"Successfully saved converted text to: {output_path}")

def main():
    if len(sys.argv) < 3:
        print("Usage: python fix_georgian_encoding.py <input_file> <output_file>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(input_path):
        print(f"Error: Input file '{input_path}' not found.")
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()

    if ext == '.docx':
        process_docx(input_path, output_path)
    elif ext == '.txt':
        process_text_file(input_path, output_path)
    else:
        print("Unsupported file format. Please use .docx or .txt")

if __name__ == "__main__":
    main()
